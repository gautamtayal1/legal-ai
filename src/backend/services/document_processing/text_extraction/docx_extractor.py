"""
Simple DOCX text extraction.
"""

from pathlib import Path
import docx
from .base import TextExtractor, ExtractionResult

class DOCXExtractor(TextExtractor):
    
    def extract(self, file_path: Path) -> ExtractionResult:
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return ExtractionResult(
                text='\n'.join(text_parts),
                metadata={
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "method": "python-docx"
                }
            )
            
        except Exception as e:
            return ExtractionResult("", error=f"DOCX extraction failed: {e}") 